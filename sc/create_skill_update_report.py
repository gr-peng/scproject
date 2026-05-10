import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
RUN_DATE = "2026-05-09"
OUT_DIR = ROOT / "outputs" / "skill-update-report"
DOCX = OUT_DIR / "daily-arxiv-agent-skill-update-report-zh.docx"

EXPERIMENTS = [
    ("skill-proof-1-agentic-ai-systems", "实验一：Agentic AI systems"),
    ("skill-proof-2-ai-for-finance", "实验二：AI for finance"),
    ("skill-proof-3-ai-for-scientific-discovery", "实验三：AI for scientific discovery"),
]

FIGURE_EXPLANATIONS = [
    ("evaluation_scorecard.png", "综合评价记分卡", "展示一次 briefing 运行的总质量分，以及检索、排序、时效性、多样性、摘要质量、链接完整性、图表生成、运行效率等分组得分。它用于快速判断整个 Agent 运行是否健康。"),
    ("quality_radar.png", "质量雷达图", "用八个维度展示同一组质量分。某一轴明显收缩时，说明该环节是当前系统短板，例如排序区分度不足或类别多样性不足。"),
    ("top_paper_matrix.png", "Top-k 论文质量矩阵", "解释每篇入选论文为什么被选中。每一行是一篇论文，每个条形分别表示相关性、创新性、引用/影响力代理、时效性、方法信号、最终得分和摘要就绪度。"),
    ("category_distribution.png", "类别分布图", "展示 top-k 论文在 arXiv primary category 上的分布，用于评价选题覆盖面，并检查是否被单一类别过度主导。"),
    ("relevance_recency.png", "相关性-时效性散点图", "横轴表示论文年龄，越靠左越新；纵轴表示相关性，越高越相关；点的颜色和大小反映最终得分。"),
    ("score_distribution.png", "评分分布图", "展示相关性、创新性、引用/影响力代理、方法信号和最终得分在不同分数区间的分布，用于判断排序是否有足够区分度。"),
    ("summary_quality.png", "摘要质量图", "展示摘要完成度、具体性、可行动性，以及 contribution、method、brief 三个字段的覆盖率。"),
    ("runtime_by_stage.png", "阶段运行时间图", "展示 Search、Ranking、Summarization、Visualization、Briefing Generation 等阶段耗时，用于检查效率和瓶颈。"),
]

SKILL_ROWS = [
    (
        "arxiv-search-skill",
        "根据用户主题或关键词检索 arXiv 论文，并返回结构化 Paper 对象。",
        "增加了受控 query expansion、metadata/source integrity 标记、可选 citation enrichment 字段，并明确要求不能编造引用量。",
        "输入包括关键词、类别、max_results、日期窗口、排序方式、query_expansion 和 enrich_citations；输出包括论文元数据、链接、类别、匹配词、可选 citation_count，以及下游评分和摘要字段。"
    ),
    (
        "paper-ranking-skill",
        "把检索结果排序成可阅读、可解释的 top-k 论文列表。",
        "从 relevance+novelty 两维评分扩展为多信号评分：相关性、创新性、引用/影响力代理、时效性、方法信号和最终得分，并增加 score_breakdown。",
        "只有外部数据源提供 citation_count 时才使用真实引用量；当前实验没有接入引用 API，因此使用明确标注的 impact_proxy_score，避免把代理分数说成真实引用量。"
    ),
    (
        "paper-summarization-skill",
        "从 abstract 中抽取贡献、方法、证据句，并生成简洁 briefing summary。",
        "增加 evidence_sentences、summary_quality_flags、fallback_used，以及 actionability/specificity 等摘要质量指标。",
        "输出 contribution_summary、method_summary、brief_summary、证据支持和质量标记；不生成 abstract 中没有依据的技术结论。"
    ),
    (
        "visualization-skill",
        "计算任务级评价指标，并生成可直接放入报告的图表。",
        "主图表格式从 SVG 改为 PNG，并把评价体系升级为检索、排序、时效性、多样性、摘要、完整性、图表生成和运行效率八组指标。",
        "生成八张 PNG 图：综合记分卡、质量雷达图、top-k 论文矩阵、类别分布、相关性-时效性散点图、评分分布、摘要质量和运行时间图。"
    ),
    (
        "briefing-generation-skill",
        "把论文数据、评分、摘要、指标和图表组装成最终 briefing 报告。",
        "Markdown 图表引用改为 PNG，指标表扩展为多维评价，summary table 增加多信号 ranking 列，并在契约中加入可选 Word 输出。",
        "输出对齐的 Markdown、CSV、JSON 和 manifest；JSON 是 source of truth，并在 Notes 中解释 citation/proxy 的回退策略。"
    ),
]


def load_experiments():
    rows = []
    for dirname, label in EXPERIMENTS:
        path = ROOT / "outputs" / dirname / f"daily-arxiv-briefing-{RUN_DATE}.json"
        data = json.loads(path.read_text(encoding="utf-8"))
        metrics = data["metrics"]
        rows.append(
            {
                "label": label,
                "dir": dirname,
                "topic": data.get("topic") or data.get("topic_raw") or label,
                "query": data.get("query", ""),
                "retrieved": int(metrics.get("retrieved_count", 0)),
                "top_k": int(metrics.get("top_k_count", data.get("top_k", 0))),
                "overall": metrics.get("overall_quality_score", 0.0),
                "ranking": metrics.get("ranking_quality_score", 0.0),
                "summary": metrics.get("summary_quality_score", 0.0),
                "figures": len(data.get("figures", [])),
                "top_paper": data.get("papers", [{}])[0].get("title", ""),
            }
        )
    return rows


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Arial"
    return heading


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    return table


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    experiments = load_experiments()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Daily arXiv Research Briefing Agent\n技能更新说明文档")
    run.bold = True
    run.font.size = Pt(20)
    subtitle = doc.add_paragraph("五个 Skills 的功能更新、PNG 可视化、多维 Ranking 评分与最新实验结果")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. 总体说明", 1)
    doc.add_paragraph(
        "本文档总结升级后的五个 Skill 组成的 Daily arXiv Research Briefing Agent。"
        "本轮主要改动包括：可视化统一改为 PNG 格式；评价体系从简单的生成成功率扩展为更贴合任务目标的质量指标；"
        "Ranking 增加 citation/proxy、recency、method signal 等多维评分；Search、Summarization 和 Briefing 的输入输出契约也同步增强。"
    )
    doc.add_paragraph(
        "注意：arXiv 元数据本身不直接提供引用量。现在 Ranking Skill 支持在外部学术数据源可用时使用真实 citation_count。"
        "当前最新实验没有接入 Semantic Scholar 或 OpenAlex 等引用 API，因此系统没有编造引用量，而是使用明确标注的 impact_proxy_score 作为影响力代理分数。"
    )

    add_heading(doc, "2. 每个 Skill 的改动与当前完整功能", 1)
    add_table(
        doc,
        ["Skill", "当前功能", "本轮主要改动", "现在的完整功能"],
        SKILL_ROWS,
    )

    add_heading(doc, "3. 可视化图片及含义", 1)
    figure_rows = [(name, title, meaning) for name, title, meaning in FIGURE_EXPLANATIONS]
    add_table(doc, ["PNG 文件", "图名", "含义"], figure_rows)

    sample_dir = ROOT / "outputs" / "skill-proof-1-agentic-ai-systems" / "figures"
    doc.add_paragraph("以下插入的是最新保留实验中 Agentic AI Systems proof run 的示例图片：")
    for filename, title_text, meaning in FIGURE_EXPLANATIONS:
        image_path = sample_dir / filename
        if image_path.exists():
            add_heading(doc, title_text, 2)
            doc.add_paragraph(meaning)
            doc.add_picture(str(image_path), width=Inches(6.4))

    add_heading(doc, "4. 当前保留并测试的实验", 1)
    exp_rows = [
        [
            exp["label"],
            exp["retrieved"],
            exp["top_k"],
            f'{exp["overall"] * 100:.1f}',
            f'{exp["ranking"] * 100:.1f}',
            f'{exp["summary"] * 100:.1f}',
            exp["figures"],
            exp["top_paper"][:80],
        ]
        for exp in experiments
    ]
    add_table(
        doc,
        ["实验", "检索数", "Top-k", "总分", "排序分", "摘要分", "图数", "Top paper"],
        exp_rows,
    )

    add_heading(doc, "5. 当前保留的输出文件", 1)
    doc.add_paragraph(
        "每个最新实验目录都包含 Markdown briefing、CSV 表格、JSON source snapshot、manifest JSON，以及 figures/ 下的八张 PNG 图。"
        "Markdown 报告现在只引用 PNG，不再引用 SVG。"
    )
    for exp in experiments:
        doc.add_paragraph(f"- outputs/{exp['dir']}: {exp['topic']}")

    add_heading(doc, "6. 后续优化建议", 1)
    recommendations = [
        "接入 Semantic Scholar 或 OpenAlex 的 citation enrichment，让 citation_score 使用真实引用量，而不是当前的 impact proxy。",
        "加入人工相关性标签，计算 Precision@5、Precision@10 和 nDCG@10，使 ranking 评估更严谨。",
        "在 abstract 信息不足时加入 PDF-level extraction，抽取更完整的 contribution、method 和 evidence。",
        "保存 previous-run JSON，实现 daily diff：新增论文、重复论文和消失论文。",
        "最终提交 StudyClawHub 前，把当前 workflow 脚本拆成五个独立模块，每个 Skill 都有独立测试和示例输出。",
    ]
    for item in recommendations:
        doc.add_paragraph(item, style=None)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
